from __future__ import annotations

import json
import logging
import os
import re
import unicodedata
from pathlib import Path
from typing import Any
from uuid import uuid4

from dotenv import load_dotenv

logger = logging.getLogger("optiguide.agent")

try:
    from docx import Document
except Exception:  # pragma: no cover - optional at runtime
    Document = None  # type: ignore

try:
    from openai import OpenAI
except Exception:  # pragma: no cover - optional at runtime
    OpenAI = None  # type: ignore


load_dotenv()

LLM_BASE_URL = os.environ.get("LLM_BASE_URL", "").rstrip("/")
LLM_MODEL = os.environ.get("LLM_MODEL", "").strip()
LLM_API_KEY = os.environ.get("LLM_API_KEY", "not-needed")
LLM_TIMEOUT = float(os.environ.get("LLM_TIMEOUT", "20"))
LLM_TEMPERATURE = float(os.environ.get("LLM_TEMPERATURE", "0.1"))
LLM_MAX_TOKENS = int(os.environ.get("LLM_MAX_TOKENS", "350"))
USE_LLM_AGENT = os.environ.get("USE_LLM_AGENT", "1").strip().lower() not in {"0", "false", "no"}


FIELD_QUESTIONS = {
    "age": "Quel age avez-vous ?",
    "head_eye_behavior": "Lorsque vous regardez a droite et a gauche, bougez-vous davantage la tete ou utilisez-vous principalement vos yeux ?",
    "postural_comfort": "Souhaitez-vous bénéficier d'un meilleur confort postural lors de la lecture ?",
    "correction_total": "Pourriez-vous me communiquer la correction de votre oeil droit et de votre oeil gauche, ainsi que l'addition si elle est mentionnée sur votre ordonnance ?",
    "near_intermediate_comfort": "Souhaitez-vous un verre qui vous permette de voir confortablement de pres et en vision intermediaire sans avoir a bouger la tete ?",
    "innovation_sensitive": "Etes-vous sensible aux solutions innovantes ?",
    "computer_usage": "Utilisez-vous principalement un ordinateur de bureau ou un ordinateur portable ?",
    "frame_type": "Quel type de montage preferez-vous (Nylor, Plastique, Metallique, Perce) ?",
    "main_need": "Quels sont vos principaux besoins concernant vos verres ?",
    "sun_exposure": "Etes-vous souvent expose au soleil dans votre quotidien ?",
    "sun_discomfort": "L'exposition au soleil vous semble-t-elle genante ou difficile a supporter ?",
    "sun_solution": "Preferez-vous un verre solaire ou une transition qui s'adapte automatiquement a la luminosite ?",
    "ocular_health": "Pouvez-vous nous indiquer votre etat de sante oculaire actuel afin d'adapter au mieux la recommandation ?",
}


SCENARIO_FLOWS = {
    "simple_foyer": [
        "age",
        "correction_total",
        "frame_type",
        "main_need",
        "sun_exposure",
        "sun_discomfort",
        "sun_solution",
        "ocular_health",
    ],
    "eyezen": [
        "age",
        "correction_total",
        "frame_type",
        "main_need",
        "sun_exposure",
        "sun_discomfort",
        "sun_solution",
        "ocular_health",
    ],
    "varilux": [
        "age",
        "head_eye_behavior",
        "postural_comfort",
        "correction_total",
        "near_intermediate_comfort",
        "innovation_sensitive",
        "computer_usage",
        "frame_type",
        "main_need",
        "sun_exposure",
        "sun_discomfort",
        "sun_solution",
        "ocular_health",
    ],
}


def _norm(text: Any) -> str:
    value = unicodedata.normalize("NFKD", str(text or ""))
    value = "".join(char for char in value if not unicodedata.combining(char))
    return re.sub(r"\s+", " ", value.lower()).strip()


def _first_number(text: Any) -> float | None:
    # Ne pas utiliser `text or ""` : la valeur 0 (ex: oeil plan) serait convertie
    # en chaine vide et perdue. On ne remplace que None par une chaine vide.
    match = re.search(r"-?\d+(?:[\.,]\d+)?", "" if text is None else str(text))
    if not match:
        return None
    return float(match.group(0).replace(",", "."))


# ---------------------------------------------------------------------------
# Normalisation des corrections optiques (format en centiemes de dioptrie).
#
# Format interne canonique = centiemes (entier) : 200 => +2.00 D, -175 => -1.75 D.
# `to_centiemes` accepte aussi l'ancien format decimal (2.00) pour compatibilite :
# une valeur avec des decimales OU |v| < 20 est interpretee comme des dioptries
# et multipliee par 100 ; sinon elle est deja en centiemes.
# L'axe est en degres et NE passe jamais par ces fonctions.
# ---------------------------------------------------------------------------
LEGACY_DIOPTER_MAX = 20  # au-dela, la valeur est forcement en centiemes


def to_centiemes(value: Any) -> int | None:
    """Convertit une correction (dioptries legacy ou centiemes) en centiemes entiers."""
    number = _first_number(value)
    if number is None:
        return None
    if number != int(number) or abs(number) < LEGACY_DIOPTER_MAX:
        return int(round(number * 100))
    return int(number)


def to_diopters(value: Any) -> float | None:
    """Convertit une correction vers des dioptries (unite metier des comparaisons)."""
    centiemes = to_centiemes(value)
    return None if centiemes is None else centiemes / 100.0


# ---------------------------------------------------------------------------
# Noms de produits canoniques (source unique) + resolution des alias/codes.
# ---------------------------------------------------------------------------
PRODUCT_XR = "Varilux XR Design"
PRODUCT_X = "Varilux X Design"
PRODUCT_S = "Varilux S Design"
PRODUCT_PHYSIO = "VX Physio 3.0"
PRODUCT_COMFORT = "Varilux Comfort 3.0"
PRODUCT_COMFORT_MAX = "Varilux Comfort 3.0 Max"
PRODUCT_LIBERTY = "Varilux Liberty"

# Alias/codes tolerts en entree -> nom canonique. La cle "rep" corrige le bug ou
# la reponse "rep" (et ses variantes de casse/espaces) ne mappait vers aucun
# produit : elle designe le Varilux XR Design.
LENS_TYPE_ALIASES = {
    "rep": PRODUCT_XR,
    "xr": PRODUCT_XR,
    "xr design": PRODUCT_XR,
    "varilux xr": PRODUCT_XR,
    "x design": PRODUCT_X,
    "s design": PRODUCT_S,
    "physio": PRODUCT_PHYSIO,
    "vx physio": PRODUCT_PHYSIO,
    "varilux physio": PRODUCT_PHYSIO,
    "comfort": PRODUCT_COMFORT,
    "varilux comfort": PRODUCT_COMFORT,
}


def canonical_lens_type(value: Any) -> str:
    """Uniformise un nom/code de verre (trim + minuscules + table d'alias).

    Ne modifie que les valeurs qui correspondent exactement a un alias connu
    (ex: "rep", "REP", " rep "), les noms complets deja canoniques restent
    inchanges.
    """
    normalized = _norm(value)
    if normalized in LENS_TYPE_ALIASES:
        return LENS_TYPE_ALIASES[normalized]
    return str(value or "").strip()


def family_from_lens_type(lens_type: Any) -> str:
    """Famille d'affichage a partir du nom de verre (gere VX Physio 3.0)."""
    normalized = _norm(lens_type)
    if "varilux" in normalized or "vx physio" in normalized or "physio" in normalized:
        return "Varilux"
    if "eyezen" in normalized or "eyzen" in normalized:
        return "Eyezen"
    return "Simple Foyer"


# ---------------------------------------------------------------------------
# Libelles des criteres (nouveaux intitules) + transition (Couleur / Gen S).
# Les valeurs techniques restent inchangees pour la compatibilite ; seuls les
# libelles affiches evoluent.
# ---------------------------------------------------------------------------
MAIN_NEED_LABELS = {
    "transparence": "Transparence",
    "lumiere_bleue": "Protection contre la lumiere bleue",
    "conduite_soir": "Securite lors de la conduite le soir",
    "rayures": "Resistance contre les rayures",
    "nettoyage": "Nettoyage facile",
}

# Choix maladies (multi-selection). Valeurs affichees = valeurs stockees.
OCULAR_CHOICES = [
    "RAS",
    "Cataracte",
    "Glaucome",
    "DMLA",
    "Pseudophaque",
    "Conjonctivite",
    "Retinopathie diabetique",
    "Arthrose",
]

TRANSITION_LABEL = "Couleur"  # ancien libelle "Transition(s)" -> "Couleur"


def display_transition(value: Any) -> str:
    """Rend une valeur de transition avec les nouveaux intitules (Gen S / Couleur)."""
    normalized = _norm(value)
    if normalized in {"transition", "photochromique", "photo"}:
        return "Gen S"
    if normalized == "transitions":
        return TRANSITION_LABEL
    return str(value or "").strip() or "-"


# ---------------------------------------------------------------------------
# Valeurs de catalogue : source unique alimentant les listes de l'interface
# Produits. Les designs Varilux/VX reutilisent les constantes PRODUCT_* pour
# eviter toute divergence. INDEX_VALUES = liste d'indices deja utilisee dans le
# projet (voir _recommend_index_from_profile).
# ---------------------------------------------------------------------------
DESIGN_VALUES = [
    "Simple Foyer",
    "Eyezen initial",
    "Eyezen actif",
    "Eyezen actif +",
    PRODUCT_LIBERTY,
    PRODUCT_COMFORT,
    PRODUCT_COMFORT_MAX,
    PRODUCT_PHYSIO,
    PRODUCT_S,
    PRODUCT_X,
    PRODUCT_XR,
]
INDEX_VALUES = ["1.50", "1.56", "1.60", "1.67"]
TREATMENT_VALUES = [
    "Crizal Easy Pro",
    "Crizal Sapphire HR",
    "Crizal Prevencia",
    "Crizal Drive",
    "Crizal Rock",
]
TRANSITION_VALUES = ["Blanc", "Gen S", "Solaire"]
GEOMETRY_VALUES = ["Regular", "Short"]


def _yes_no(text: Any) -> bool | None:
    normalized = _norm(text)
    words = set(normalized.split())
    if words & {"oui", "yes", "ok", "daccord", "exact"}:
        return True
    if words & {"non", "no", "jamais", "pas"} or normalized.startswith("pas "):
        return False
    return None


def _profile_bool(value: Any) -> bool | None:
    if isinstance(value, bool):
        return value
    normalized = _norm(value)
    if normalized in {"true", "oui", "yes", "1"}:
        return True
    if normalized in {"false", "non", "no", "0"}:
        return False
    return None


def _family_from_age(age: int | float | None) -> str | None:
    if age is None:
        return None
    if age <= 30:
        return "simple_foyer"
    if age <= 41:
        return "eyezen"
    return "varilux"


def _ocular_text(profile: dict[str, Any]) -> str:
    """Texte normalise des maladies (accepte une chaine ou une liste)."""
    value = profile.get("ocular_health")
    if isinstance(value, (list, tuple, set)):
        return _norm(" ".join(str(item) for item in value))
    return _norm(value or "")


def _extract_eye_values(text: str, eye: str) -> dict[str, int | None]:
    """Extrait sphere / cylindre / addition (en centiemes) pour un oeil (od/og).

    L'axe est ignore ici : il reste en degres et n'entre pas dans le calcul de
    correction. Renvoie None pour un parametre absent.
    """
    normalized = _norm(text)
    match = re.search(rf"\b{eye}\b(?P<section>.*?)(?=\bod\b|\bog\b|$)", normalized)
    if not match:
        return {"sph": None, "cyl": None, "add": None}
    section = match.group("section")

    def grab(keys: list[str]) -> int | None:
        for key in keys:
            found = re.search(rf"\b{key}\b\s*([+-]?\d+(?:[\.,]\d+)?)", section)
            if found:
                return to_centiemes(found.group(1))
        return None

    sphere = grab(["sph", "sphere", "shp"])
    if sphere is None:
        # Pas de mot-cle "sph" : premier nombre = sphere par convention.
        first = re.search(r"([+-]?\d+(?:[\.,]\d+)?)", section)
        if first:
            sphere = to_centiemes(first.group(1))
    return {"sph": sphere, "cyl": grab(["cyl", "cylindre"]), "add": grab(["add", "addition"])}


def _eye_correction(values: dict[str, int | None]) -> int:
    """Correction globale d'un oeil (centiemes) = sphere + cylindre + addition."""
    return sum(int(values.get(component) or 0) for component in ("sph", "cyl", "add"))


def _recommend_index_from_profile(profile: dict[str, Any]) -> str:
    if profile.get("frame_type") == "perce":
        return "1.60"
    correction = to_centiemes(profile.get("correction_total"))
    if correction is None:
        return "1.50"
    if correction <= 200:
        return "1.50"
    if correction <= 500:
        return "1.56"
    if correction <= 800:
        return "1.60"
    return "1.67"


def _recommend_treatment_from_profile(profile: dict[str, Any]) -> tuple[str, str]:
    family = _family_from_age(profile.get("age") if isinstance(profile.get("age"), (int, float)) else None)
    ocular = _ocular_text(profile)
    prevencia_terms = ["pseudophaque", "glaucome", "cataracte", "dmla", "conjonctivite", "retinopathie"]
    if family == "eyezen":
        prevencia_terms.append("arthrose")
    if any(term in ocular for term in prevencia_terms):
        return "Crizal Prevencia", "Sante oculaire signalee => Crizal Prevencia prioritaire."
    mapping = {
        "transparence": ("Crizal Sapphire HR", "Besoin transparence => Crizal Sapphire HR."),
        "lumiere_bleue": ("Crizal Prevencia", "Protection lumiere bleue => Crizal Prevencia."),
        "conduite_soir": ("Crizal Drive", "Conduite de nuit => Crizal Drive."),
        "rayures": ("Crizal Rock", "Resistance aux rayures => Crizal Rock."),
        "nettoyage": ("Crizal Easy Pro", "Facilite de nettoyage => Crizal Easy Pro."),
    }
    return mapping.get(str(profile.get("main_need")), ("Crizal Easy Pro", "Traitement par defaut => Crizal Easy Pro."))


def _recommend_transition_from_profile(profile: dict[str, Any]) -> tuple[str, str]:
    if _profile_bool(profile.get("sun_exposure")) is False:
        return "Blanc", "Pas d'exposition frequente au soleil => Blanc."
    if _profile_bool(profile.get("sun_discomfort")) is False:
        return "Blanc", "Exposition solaire non genante => Blanc."
    solution = _norm(
        profile.get("sun_solution")
        or profile.get("transition_preference")
    )
    if "solaire" in solution:
        return "Solaire", "Preference solaire => Solaire."
    if "photo" in solution or "transition" in solution or "gen s" in solution:
        return "Gen S", "Preference Gen S (photochromique) => Gen S."
    return "Blanc", "Couleur par defaut => Blanc."


def _recommend_varilux_type(profile: dict[str, Any]) -> tuple[str, list[str]]:
    reasons: list[str] = []
    behavior = profile.get("head_eye_behavior")
    if behavior == "head":
        lens_type = PRODUCT_LIBERTY
        reasons.append(f"Client bouge davantage la tete => {PRODUCT_LIBERTY}.")
    elif behavior == "eyes":
        lens_type = PRODUCT_COMFORT
        reasons.append(f"Client utilise davantage les yeux => {PRODUCT_COMFORT}.")
        if _profile_bool(profile.get("postural_comfort")) is True:
            lens_type = PRODUCT_COMFORT_MAX
            reasons.append(f"Question complementaire {PRODUCT_COMFORT} : confort postural souhaite => {PRODUCT_COMFORT_MAX}.")
        elif _profile_bool(profile.get("postural_comfort")) is False:
            reasons.append(f"Question complementaire {PRODUCT_COMFORT} : confort postural non souhaite => conserver {PRODUCT_COMFORT}.")
    else:
        lens_type = PRODUCT_LIBERTY
        reasons.append(f"Valeur par defaut Varilux => {PRODUCT_LIBERTY}.")

    # ADD comparee en dioptries (accepte centiemes ou ancien format decimal).
    add_power = to_diopters(profile.get("add_power"))
    if add_power is not None and add_power > 2.25:
        lens_type = PRODUCT_S
        reasons.append(f"ADD > 2.25 => {PRODUCT_S}.")
    if _profile_bool(profile.get("near_intermediate_comfort")) is True:
        lens_type = PRODUCT_X
        reasons.append(f"Vision proche/intermediaire sans bouger la tete souhaitee => {PRODUCT_X}.")
    if _profile_bool(profile.get("innovation_sensitive")) is True:
        lens_type = PRODUCT_XR
        reasons.append(f"Client sensible aux solutions innovantes => {PRODUCT_XR}.")

    # Regle VX Physio 3.0 (prioritaire) : difference des CORRECTIONS TOTALES OD/OG
    # (sphere + cylindre + addition par oeil) strictement superieure a 2.00 D.
    correction_od = to_centiemes(profile.get("correction_od"))
    correction_og = to_centiemes(profile.get("correction_og"))
    if correction_od is not None and correction_og is not None and abs(correction_od - correction_og) > 200:
        lens_type = PRODUCT_PHYSIO
        reasons.append(f"Difference de correction totale OD/OG > 2.00 D => {PRODUCT_PHYSIO}.")
    return lens_type, reasons


def _recommend_geometry_from_profile(profile: dict[str, Any], family: str | None) -> tuple[str, str]:
    if family != "varilux":
        return "", ""
    ocular = _ocular_text(profile)
    if "arthrose" in ocular:
        return "Short", "Arthrose => geometrie Short."
    if "pseudophaque" in ocular:
        return "Short", "Pseudophaque => geometrie Short."
    if profile.get("computer_usage") == "high":
        return "Short", "Ordinateur portable => Short."
    return "Regular", "Ordinateur de bureau => Regular."


# ---------------------------------------------------------------------------
# Justifications concises orientees "produit" (affichage). Elles expliquent
# POURQUOI la valeur est recommandee sans reciter le detail des questions.
# La trace detaillee reste disponible via "rationale" / "applied_rules".
# ---------------------------------------------------------------------------
def _justify_type(lens_type: str) -> str:
    normalized = _norm(lens_type)
    if "physio" in normalized:
        return "Progressif recommande pour equilibrer une difference de correction marquee entre les deux yeux."
    if "xr design" in normalized:
        return "Progressif haut de gamme, la solution la plus innovante et la plus fluide."
    if "x design" in normalized:
        return "Progressif offrant une large vision de pres et en vision intermediaire."
    if "s design" in normalized:
        return "Progressif adapte a une addition elevee."
    if "comfort" in normalized:
        return "Progressif offrant un large confort de lecture au quotidien."
    if "liberty" in normalized:
        return "Progressif adapte a votre usage courant."
    if "eyezen" in normalized:
        return "Verre concu pour soulager la vision de pres et la fatigue visuelle."
    if "simple" in normalized:
        return "Verre unifocal (simple foyer) adapte a votre correction."
    return f"{lens_type} adapte a votre profil."


def _justify_index(index: str, profile: dict[str, Any]) -> str:
    if profile.get("frame_type") == "perce":
        return f"Indice {index} requis par un montage perce."
    return f"Indice {index} choisi selon la puissance de votre correction."


def _justify_treatment(treatment: str) -> str:
    normalized = _norm(treatment)
    if "prevencia" in normalized:
        return "Protege vos yeux (lumiere bleue et sante oculaire)."
    if "sapphire" in normalized:
        return "Transparence maximale et reflets reduits."
    if "drive" in normalized:
        return "Optimise la vision et le confort pour la conduite de nuit."
    if "rock" in normalized:
        return "Resistance renforcee aux rayures."
    if "easy" in normalized:
        return "Entretien facile et nettoyage rapide."
    return "Traitement adapte a vos besoins."


def _justify_transition(transition: str) -> str:
    normalized = _norm(transition)
    if "gen s" in normalized or "transition" in normalized or "photo" in normalized:
        return "Verre photochromique Gen S qui s'adapte automatiquement a la luminosite."
    if "solaire" in normalized:
        return "Verre solaire adapte a une forte exposition au soleil."
    return "Verre blanc : pas de besoin de protection solaire particuliere."


def _justify_geometry(geometry: str, family: str | None) -> str:
    if family != "varilux" or not geometry:
        return "Non applicable pour ce type de verre."
    if _norm(geometry) == "short":
        return "Geometrie Short adaptee a un usage rapproche et a plus de confort."
    return "Geometrie Regular adaptee a un usage standard."


def build_document_recommendation(profile: dict[str, Any], family: str | None) -> dict[str, Any]:
    rationale: list[str] = []
    if family == "simple_foyer":
        lens_type = "simple foyer"
        rationale.append("Age <= 30 ans => simple foyer.")
    elif family == "eyezen":
        age = int(profile.get("age") or 0)
        if 31 <= age <= 34:
            lens_type = "Eyezen initial"
            rationale.append("Age 31-34 ans => Eyezen initial.")
        elif 35 <= age <= 37:
            lens_type = "Eyezen actif"
            rationale.append("Age 35-37 ans => Eyezen actif.")
        else:
            lens_type = "Eyezen actif +"
            rationale.append("Age 38-41 ans => Eyezen actif +.")
    else:
        lens_type, rationale = _recommend_varilux_type(profile)

    lens_type = canonical_lens_type(lens_type)
    # Justification propre au Design (les raisons qui ont produit le type de verre).
    type_reason = " ".join(rationale)

    index = _recommend_index_from_profile(profile)
    # La gamme Varilux S Design n'existe pas en indice 1.56 : on monte au 1.60.
    s_design_index_bumped = "s design" in _norm(lens_type) and index == "1.56"
    if s_design_index_bumped:
        index = "1.60"
    if profile.get("frame_type") == "perce":
        index_reason = "Montage perce => indice 1.60 prioritaire."
    elif s_design_index_bumped:
        index_reason = "Varilux S Design indisponible en 1.56 => indice 1.60."
    else:
        index_reason = f"Indice calcule selon la puissance maximale OD/OG => {index}."
    rationale.append(index_reason)
    treatment, treatment_reason = _recommend_treatment_from_profile(profile)
    transition, transition_reason = _recommend_transition_from_profile(profile)
    geometry, geometry_reason = _recommend_geometry_from_profile(profile, family)
    rationale.extend([treatment_reason, transition_reason])
    if geometry_reason:
        rationale.append(geometry_reason)
    # Justifications concises par critere (affichage) ; la trace detaillee reste
    # dans "rationale". type_reason/index_reason ci-dessus alimentent "rationale".
    _ = type_reason  # conserve pour la trace detaillee (rationale)
    justifications = {
        "type": _justify_type(lens_type),
        "index": _justify_index(index, profile),
        "treatment": _justify_treatment(treatment),
        "transition": _justify_transition(transition),
    }
    if geometry:
        justifications["geometrie"] = _justify_geometry(geometry, family)
    reco = {
        "lens_type": lens_type,
        "index": index,
        "indice": index,
        "treatment": treatment,
        "traitement": treatment,
        "transition": transition,
        "confidence": 0.95,
        "rationale": rationale,
        "applied_rules": rationale,
        "justifications": justifications,
    }
    if geometry:
        reco["geometrie"] = geometry
    return build_price_optimization(reco)


def _downgrade_type(lens_type: str) -> str | None:
    normalized = _norm(lens_type)
    if "simple" in normalized or "eyezen" in normalized:
        return None
    if "xr" in normalized:
        return PRODUCT_X
    if "x design" in normalized:
        return PRODUCT_S
    if "s design" in normalized:
        return PRODUCT_PHYSIO
    if "physio" in normalized:
        return PRODUCT_COMFORT
    if "comfort" in normalized and "max" in normalized:
        return PRODUCT_COMFORT
    if "comfort" in normalized:
        return PRODUCT_LIBERTY
    return None


def _downgrade_treatment(treatment: str) -> str | None:
    normalized = _norm(treatment).replace(" ", "")
    if normalized in {"crizalprevencia", "crizaldrive", "crizalsapphirehr"}:
        return "Crizal Rock"
    if normalized == "crizalrock":
        return "Crizal Easy Pro"
    return None


def _downgrade_transition(transition: str) -> str | None:
    normalized = _norm(transition)
    if "polaris" in normalized:
        return "Gen S"
    if "gen s" in normalized or "transition" in normalized or "photo" in normalized or "solaire" in normalized:
        return "Blanc"
    return None


def _available_downgrades(reco: dict[str, Any]) -> dict[str, str]:
    options: dict[str, str] = {}
    next_type = _downgrade_type(str(reco.get("lens_type", "")))
    next_treatment = _downgrade_treatment(str(reco.get("treatment", reco.get("traitement", ""))))
    next_transition = _downgrade_transition(str(reco.get("transition", "")))
    if next_type:
        options["type"] = next_type
    if next_treatment:
        options["treatment"] = next_treatment
    if next_transition:
        options["transition"] = next_transition
    return options


def build_price_optimization(reco: dict[str, Any]) -> dict[str, Any]:
    enriched = dict(reco)
    attempts = int(enriched.get("priceOptimization", {}).get("attempts", enriched.get("downgradeAttempts", 0)) or 0)
    available = _available_downgrades(enriched) if attempts < 2 else {}
    enriched["priceOptimization"] = {
        "enabled": bool(available),
        "attempts": attempts,
        "maxAttempts": 2,
        "availableDowngrades": available,
        "message": (
            "Cette recommandation peut être optimisée afin de réduire le prix. "
            "L'indice et la géométrie sont conservés."
            if available
            else "Aucune réduction supplémentaire disponible ou limite de deux tentatives atteinte."
        ),
    }
    enriched["downgradeAttempts"] = attempts
    return enriched


def apply_price_downgrade(reco: dict[str, Any], field: str) -> dict[str, Any]:
    updated = dict(reco)
    optimization = updated.get("priceOptimization", {}) if isinstance(updated.get("priceOptimization"), dict) else {}
    attempts = int(optimization.get("attempts", updated.get("downgradeAttempts", 0)) or 0)
    if attempts >= 2:
        return build_price_optimization(updated)

    available = optimization.get("availableDowngrades", {}) if isinstance(optimization.get("availableDowngrades"), dict) else _available_downgrades(updated)
    next_value = available.get(field)
    if not next_value:
        return build_price_optimization(updated)

    before_index = updated.get("index", updated.get("indice", ""))
    before_geometry = updated.get("geometrie", "")
    if field == "type":
        updated["lens_type"] = next_value
    elif field == "treatment":
        updated["treatment"] = next_value
        updated["traitement"] = next_value
    elif field == "transition":
        updated["transition"] = next_value

    if before_index:
        updated["index"] = before_index
        updated["indice"] = before_index
    if before_geometry:
        updated["geometrie"] = before_geometry

    rationale = list(updated.get("rationale", []) or [])
    downgrade_note = f"Optimisation prix appliquée sur {field}; indice et géométrie conservés."
    rationale.append(downgrade_note)
    updated["rationale"] = rationale
    # Garde la justification concise coherente avec la nouvelle valeur.
    justifications = dict(updated.get("justifications", {}) or {})
    if field == "type":
        justifications["type"] = _justify_type(str(updated.get("lens_type", "")))
    elif field == "treatment":
        justifications["treatment"] = _justify_treatment(str(updated.get("treatment", "")))
    elif field == "transition":
        justifications["transition"] = _justify_transition(str(updated.get("transition", "")))
    updated["justifications"] = justifications
    updated["downgradeAttempts"] = attempts + 1
    updated["priceOptimization"] = {"attempts": attempts + 1}
    return build_price_optimization(updated)


class LocalDocRAG:
    """Small local RAG over data/docs/*.docx for tool-grounded agent context."""

    def __init__(self, docs_dir: str | Path = "data/docs"):
        self.docs_dir = Path(docs_dir)
        self.chunks: list[dict[str, str]] = []
        self._load()

    def _load(self) -> None:
        self.chunks = []
        if Document is None or not self.docs_dir.exists():
            return
        for path in sorted(self.docs_dir.glob("*.docx")):
            if path.name.startswith("~$"):
                continue
            try:
                doc = Document(path)
            except Exception:
                logger.warning("Echec de lecture du scenario %s, document ignore.", path.name, exc_info=True)
                continue
            lines: list[str] = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    lines.append(text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip().replace("\n", " | ") for cell in row.cells if cell.text.strip()]
                    if cells:
                        lines.append(" | ".join(cells))
            for index in range(0, len(lines), 8):
                text = "\n".join(lines[index : index + 8])
                self.chunks.append({"source": path.name, "text": text})

    def search(self, query: str, limit: int = 4) -> list[dict[str, str]]:
        normalized_query = _norm(query)
        terms = [term for term in re.split(r"\W+", normalized_query) if len(term) > 2]
        scored: list[tuple[int, dict[str, str]]] = []
        for chunk in self.chunks:
            haystack = _norm(chunk["source"] + " " + chunk["text"])
            score = sum(3 if term in _norm(chunk["source"]) else 1 for term in terms if term in haystack)
            if score:
                scored.append((score, chunk))
        scored.sort(key=lambda item: item[0], reverse=True)
        return [chunk for _, chunk in scored[:limit]]


AGENT_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "rag_search",
            "description": "Recherche dans les documents de scenario Essilor pour ancrer la prochaine decision.",
            "parameters": {
                "type": "object",
                "properties": {"query": {"type": "string"}, "limit": {"type": "integer", "default": 4}},
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "update_profile",
            "description": "Met a jour les champs du profil optique extraits de la reponse client.",
            "parameters": {
                "type": "object",
                "properties": {"updates": {"type": "object", "additionalProperties": True}},
                "required": ["updates"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "next_question",
            "description": "Retourne la prochaine question officielle selon le scenario et les champs deja collectes.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "build_recommendation",
            "description": "Construit la recommandation officielle depuis le profil actuel.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
]


class EssilorAgent:
    def __init__(self, rag: LocalDocRAG | None = None):
        self.session_id = str(uuid4())
        self.profile: dict[str, Any] = {}
        self.current_field: str | None = "age"
        self.family: str | None = None
        self.last_recommendation: dict[str, Any] | None = None
        # RAG is read-only over the scenario docs: it can be shared across
        # sessions. Per-client mutable state stays on the instance.
        self.rag = rag if rag is not None else LocalDocRAG()
        self.tool_trace: list[dict[str, Any]] = []
        self._locally_set_fields: set[str] = set()
        self.llm = None
        if USE_LLM_AGENT and LLM_BASE_URL and LLM_MODEL and OpenAI is not None:
            self.llm = OpenAI(base_url=LLM_BASE_URL, api_key=LLM_API_KEY, timeout=LLM_TIMEOUT)

    def reset(self) -> tuple[str, dict[str, Any] | None]:
        self.session_id = str(uuid4())
        self.profile = {}
        self.current_field = "age"
        self.family = None
        self.last_recommendation = None
        self.tool_trace = []
        self._locally_set_fields = set()
        return FIELD_QUESTIONS["age"], None

    def chat(self, user_message: str) -> tuple[str, dict[str, Any] | None]:
        local_updates = self._infer_profile_updates(user_message)
        self.profile.update(local_updates)
        # Local extraction is authoritative for this turn: the LLM may enrich
        # other fields but must not overwrite what we just parsed deterministically.
        self._locally_set_fields = set(local_updates)
        self.family = _family_from_age(self.profile.get("age") if isinstance(self.profile.get("age"), (int, float)) else None)

        llm_text = self._run_llm_agent_turn(user_message, local_updates) if self.llm is not None else ""

        next_field = self._next_scenario_field()
        if next_field:
            self.current_field = next_field
            return FIELD_QUESTIONS[next_field], None
        self.current_field = None
        self.last_recommendation = build_document_recommendation(self.profile, self.family)
        return llm_text or "Recommandation finalisee.", self.last_recommendation

    def _run_llm_agent_turn(self, user_message: str, local_updates: dict[str, Any]) -> str:
        prompt = (
            "Tu es OptiGuide, un agent opticien Essilor. Tu dois utiliser les tools disponibles "
            "pour chercher dans le RAG, mettre a jour le profil, determiner la prochaine question "
            "ou construire la recommandation. Ne fabrique pas de regle hors documents. "
            "Respecte les regles importantes: age route la famille; Varilux Comfort 3.0 + confort postural Oui => Varilux Comfort 3.0 Max, Non => conserver Varilux Comfort 3.0; "
            "si |correction totale OD - correction totale OG| > 2.00 D => VX Physio 3.0 (correction totale = sphere+cylindre+addition par oeil, strictement > 2.00 D); "
            "geometrie Varilux: bureau Regular, portable Short, arthrose/pseudophaque Short; axe et ADD ne choisissent pas l'indice. "
            "Corrections en centiemes (200 = 2.00 D). Noms de produits exacts: VX Physio 3.0, Comfort 3.0, Varilux XR Design, Gen S; transition affichee 'Couleur' avec valeurs Blanc/Gen S/Solaire. "
            "Pas d'addition pour Eyezen ni Simple Foyer. Les maladies peuvent etre multiples."
        )
        messages: list[dict[str, Any]] = [
            {"role": "system", "content": prompt},
            {
                "role": "user",
                "content": json.dumps(
                    {
                        "message_client": user_message,
                        "champ_actuel": self.current_field,
                        "profil_actuel": self.profile,
                        "extraction_locale": local_updates,
                        "prochaine_question_locale": self._next_scenario_field(),
                    },
                    ensure_ascii=False,
                ),
            },
        ]
        try:
            for _ in range(4):
                response = self.llm.chat.completions.create(
                    model=LLM_MODEL,
                    messages=messages,
                    tools=AGENT_TOOLS,
                    tool_choice="auto",
                    temperature=LLM_TEMPERATURE,
                    max_tokens=LLM_MAX_TOKENS,
                )
                message = response.choices[0].message
                if not message.tool_calls:
                    return message.content or ""
                messages.append(message.model_dump(exclude_none=True))
                for tool_call in message.tool_calls:
                    result = self._execute_agent_tool(tool_call.function.name, json.loads(tool_call.function.arguments or "{}"))
                    messages.append({"role": "tool", "tool_call_id": tool_call.id, "content": json.dumps(result, ensure_ascii=False, default=str)})
            return ""
        except Exception as exc:
            self.tool_trace.append({"tool": "llm_agent_error", "error": str(exc)})
            logger.warning("Tour LLM en echec, repli sur la logique locale: %s", exc, exc_info=True)
            return ""

    def _execute_agent_tool(self, name: str, args: dict[str, Any]) -> Any:
        if name == "rag_search":
            result = self.rag.search(str(args.get("query", "")), int(args.get("limit", 4) or 4))
        elif name == "update_profile":
            updates = args.get("updates", {}) if isinstance(args.get("updates"), dict) else {}
            allowed = {
                "age",
                "head_eye_behavior",
                "postural_comfort",
                "correction_total",
                "add_power",
                "near_intermediate_comfort",
                "innovation_sensitive",
                "computer_usage",
                "frame_type",
                "main_need",
                "sun_exposure",
                "sun_discomfort",
                "sun_solution",
                "transition_preference",
                "ocular_health",
            }
            clean_updates = {
                key: value
                for key, value in updates.items()
                if key in allowed and value not in (None, "") and key not in self._locally_set_fields
            }
            rejected = sorted(set(updates) & self._locally_set_fields)
            if rejected:
                logger.info("Updates LLM ignores (extraction locale prioritaire): %s", rejected)
            self.profile.update(clean_updates)
            self.family = _family_from_age(self.profile.get("age") if isinstance(self.profile.get("age"), (int, float)) else None)
            result = {"profile": self.profile, "accepted_updates": clean_updates, "rejected_fields": rejected}
        elif name == "next_question":
            field = self._next_scenario_field()
            result = {"field": field, "question": FIELD_QUESTIONS.get(field or "", "")}
        elif name == "build_recommendation":
            result = build_document_recommendation(self.profile, self.family)
        else:
            result = {"error": f"Tool inconnu: {name}"}
        self.tool_trace.append({"tool": name, "args": args, "result": result})
        return result

    def _scenario_flow(self) -> list[str]:
        self.family = _family_from_age(self.profile.get("age") if isinstance(self.profile.get("age"), (int, float)) else None)
        if self.family is None:
            return ["age"]
        flow = list(SCENARIO_FLOWS[self.family])
        if self.family == "varilux" and self.profile.get("head_eye_behavior") != "eyes":
            flow = [field for field in flow if field != "postural_comfort"]
        sun_exposure = _profile_bool(self.profile.get("sun_exposure"))
        sun_discomfort = _profile_bool(self.profile.get("sun_discomfort"))
        if sun_exposure is not True:
            flow = [field for field in flow if field not in {"sun_discomfort", "sun_solution"}]
        elif sun_discomfort is not True:
            flow = [field for field in flow if field != "sun_solution"]
        return flow

    def _field_answered(self, field: str) -> bool:
        if field in {"sun_exposure", "sun_discomfort", "postural_comfort", "near_intermediate_comfort", "innovation_sensitive"}:
            return _profile_bool(self.profile.get(field)) is not None
        value = self.profile.get(field)
        normalized = _norm(value)
        if field == "frame_type":
            return normalized in {"perce", "nylor", "plastique", "metallique"}
        if field == "main_need":
            return normalized in {"transparence", "lumiere_bleue", "conduite_soir", "rayures", "nettoyage"}
        if field == "computer_usage":
            return normalized in {"high", "low"}
        if field == "sun_solution":
            return normalized in {"transition", "solaire"}
        if field == "correction_total":
            return _first_number(value) is not None
        if field == "ocular_health":
            # Selection multiple : repondu des qu'une liste (meme ["RAS"]) est posee.
            if isinstance(value, (list, tuple)):
                return len(value) > 0
            return value not in (None, "")
        return value not in (None, "")

    def _next_scenario_field(self) -> str | None:
        for field in self._scenario_flow():
            if not self._field_answered(field):
                return field
        return None

    def progress(self) -> dict[str, Any]:
        """Avancement du parcours pour l'aperçu live (barre + questions restantes).

        Le flow depend de la famille (deduite de l'age) et des branches
        conditionnelles, donc le total peut evoluer au fil des reponses.
        """
        flow = self._scenario_flow()
        answered = [field for field in flow if self._field_answered(field)]
        remaining = [field for field in flow if not self._field_answered(field)]
        total = len(flow)
        done = self.last_recommendation is not None
        return {
            "answered": total if done else len(answered),
            "total": total,
            "ratio": 1.0 if done else (len(answered) / total if total else 0.0),
            "remaining_fields": [] if done else remaining,
            "remaining_questions": [] if done else [FIELD_QUESTIONS.get(field, field) for field in remaining],
            "done": done,
        }

    def _infer_profile_updates(self, user_message: str) -> dict[str, Any]:
        normalized = _norm(user_message)
        updates: dict[str, Any] = {}
        field = self.current_field
        if field == "age":
            number = _first_number(user_message)
            if number is not None and 1 <= number <= 100:
                updates["age"] = int(number)
            return updates
        if field == "correction_total":
            od = _extract_eye_values(user_message, "od")
            og = _extract_eye_values(user_message, "og")
            # Addition globale eventuelle si non renseignee par oeil.
            global_add = re.search(r"\badd(?:ition)?\s*([+-]?\d+(?:[\.,]\d+)?)", normalized)
            global_add_c = to_centiemes(global_add.group(1)) if global_add else None
            if od.get("add") is None:
                od["add"] = global_add_c
            if og.get("add") is None:
                og["add"] = global_add_c

            if od.get("sph") is not None or og.get("sph") is not None:
                # Sphere par oeil (centiemes) : sert a la regle VX Physio 3.0.
                if od.get("sph") is not None:
                    updates["sphere_od"] = od["sph"]
                if og.get("sph") is not None:
                    updates["sphere_og"] = og["sph"]
                if od.get("cyl") is not None:
                    updates["cyl_od"] = od["cyl"]
                if og.get("cyl") is not None:
                    updates["cyl_og"] = og["cyl"]
                add_values = [value for value in (od.get("add"), og.get("add")) if value is not None]
                if add_values:
                    updates["add_od"] = od.get("add") or 0
                    updates["add_og"] = og.get("add") or 0
                    updates["add_power"] = max(add_values)  # centiemes
                # Correction globale par oeil = sphere + cylindre + addition.
                correction_od = _eye_correction(od)
                correction_og = _eye_correction(og)
                updates["correction_od"] = correction_od
                updates["correction_og"] = correction_og
                # Indice base sur la correction max en valeur absolue (centiemes).
                updates["correction_total"] = max(abs(correction_od), abs(correction_og))
            else:
                number = to_centiemes(user_message)
                if number is not None:
                    updates["correction_total"] = abs(number)
            return updates
        if field in {"sun_exposure", "sun_discomfort", "postural_comfort", "near_intermediate_comfort", "innovation_sensitive"}:
            answer = _yes_no(user_message)
            if answer is not None:
                updates[field] = answer
                if field == "sun_exposure" and answer is False:
                    updates["transition_preference"] = "blanc"
                if field == "sun_discomfort" and answer is False:
                    updates["transition_preference"] = "blanc"
            return updates
        if field == "sun_solution":
            if "photo" in normalized or "transition" in normalized:
                updates["sun_solution"] = "transition"
                updates["transition_preference"] = "transition"
            elif "solaire" in normalized or "soleil" in normalized:
                updates["sun_solution"] = "solaire"
                updates["transition_preference"] = "solaire"
            return updates
        if field == "frame_type":
            if "perce" in normalized:
                updates["frame_type"] = "perce"
            elif "nylor" in normalized:
                updates["frame_type"] = "nylor"
            elif "plastique" in normalized:
                updates["frame_type"] = "plastique"
            elif "metal" in normalized or "metallique" in normalized:
                updates["frame_type"] = "metallique"
            return updates
        if field == "main_need":
            if "transpar" in normalized:
                updates["main_need"] = "transparence"
            elif "bleu" in normalized or "ecran" in normalized:
                updates["main_need"] = "lumiere_bleue"
            elif "conduite" in normalized or "nuit" in normalized:
                updates["main_need"] = "conduite_soir"
            elif "rayure" in normalized or "resistan" in normalized:
                updates["main_need"] = "rayures"
            elif "nettoy" in normalized or "anti reflet" in normalized or "antireflet" in normalized:
                updates["main_need"] = "nettoyage"
            return updates
        if field == "head_eye_behavior":
            has_head = "tete" in normalized
            has_eyes = "yeux" in normalized or "oeil" in normalized
            if has_head and not has_eyes:
                updates["head_eye_behavior"] = "head"
            elif has_eyes and not has_head:
                updates["head_eye_behavior"] = "eyes"
            return updates
        if field == "computer_usage":
            if "portable" in normalized:
                updates["computer_usage"] = "high"
            elif "bureau" in normalized:
                updates["computer_usage"] = "low"
            return updates
        if field == "ocular_health":
            # Selection multiple : plusieurs maladies separees par des virgules /
            # points-virgules / slash. On stocke toujours une liste ("maladies").
            parts = [part.strip() for part in re.split(r"[;,/]", user_message) if part.strip()]
            if not parts or normalized in {"ras", "rien", "aucun", "non"}:
                updates["ocular_health"] = ["RAS"]
            else:
                updates["ocular_health"] = parts
            return updates
        return updates
